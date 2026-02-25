import io
from docx import Document
from docx.shared import Pt, RGBColor


def make_docx(text):
    """Convert plain-text ATS resume into a styled .docx file."""
    doc = Document()

    # Set default Normal style
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    for line in text.strip().splitlines():
        s = line.strip()
        if not s:
            # Blank line = spacer paragraph
            doc.add_paragraph()
        elif s.isupper() and len(s) < 60:
            # ALL CAPS short line = section heading
            h = doc.add_heading(s, level=1)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                h.runs[0].font.size = Pt(13)
        elif s.startswith("-"):
            # Bullet point
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(s[1:].strip())
            run.font.size = Pt(11)
        elif "|" in s and s.count("|") >= 1 and len(s) < 120:
            # Company | Title | Dates line — bold
            p = doc.add_paragraph()
            run = p.add_run(s)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
        else:
            # Normal paragraph
            p = doc.add_paragraph(s)
            p.runs[0].font.size = Pt(11) if p.runs else None

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
